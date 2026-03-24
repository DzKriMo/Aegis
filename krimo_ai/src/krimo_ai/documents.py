from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Union


@dataclass
class DocumentResult:
    success: bool
    content: str = ""
    metadata: Dict[str, Any] = None
    error: Optional[str] = None
    pages: int = 0

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "content": self.content,
            "metadata": self.metadata,
            "error": self.error,
            "pages": self.pages,
        }


class TextDocumentProcessor:
    def read(self, path: Union[str, Path]) -> DocumentResult:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return DocumentResult(
                success=True,
                content=content,
                metadata={"path": str(path), "size": os.path.getsize(path)},
            )
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def write(self, path: Union[str, Path], content: str) -> DocumentResult:
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            return DocumentResult(
                success=True,
                content=content,
                metadata={"path": str(path), "bytes_written": len(content.encode("utf-8"))},
            )
        except Exception as e:
            return DocumentResult(success=False, error=str(e))


class MarkdownProcessor:
    def parse(self, content: str) -> Dict[str, Any]:
        sections = {}
        current_section = "introduction"
        current_content = []

        for line in content.split("\n"):
            if line.startswith("# "):
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = "h1_" + line[2:].strip().lower().replace(" ", "_")
                current_content = []
            elif line.startswith("## "):
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = "h2_" + line[3:].strip().lower().replace(" ", "_")
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def to_html(self, markdown: str) -> str:
        html = markdown
        html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
        html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
        html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)
        html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
        html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)
        html = re.sub(r"`(.+?)`", r"<code>\1</code>", html)
        html = re.sub(r"\n\n", r"</p>\n\n<p>", html)
        return f"<p>{html}</p>"


class PDFProcessor:
    def read(self, path: Union[str, Path]) -> DocumentResult:
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            pages = len(reader.pages)

            content_parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                content_parts.append(f"--- Page {i + 1} ---\n{text}")

            return DocumentResult(
                success=True,
                content="\n\n".join(content_parts),
                metadata={
                    "path": str(path),
                    "pages": pages,
                    "size": os.path.getsize(path),
                },
                pages=pages,
            )
        except ImportError:
            return DocumentResult(success=False, error="pypdf not installed. Run: pip install pypdf")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def extract_metadata(self, path: Union[str, Path]) -> Dict[str, Any]:
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            return {
                "pages": len(reader.pages),
                "metadata": reader.metadata,
            }
        except Exception:
            return {}


class WordProcessor:
    def read(self, path: Union[str, Path]) -> DocumentResult:
        try:
            from docx import Document
            doc = Document(path)
            content = "\n".join([p.text for p in doc.paragraphs])
            return DocumentResult(
                success=True,
                content=content,
                metadata={
                    "path": str(path),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                },
            )
        except ImportError:
            return DocumentResult(success=False, error="python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def write(self, path: Union[str, Path], content: str, title: str = "") -> DocumentResult:
        try:
            from docx import Document
            doc = Document()
            if title:
                doc.add_heading(title, 0)
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.save(path)
            return DocumentResult(success=True, metadata={"path": str(path)})
        except ImportError:
            return DocumentResult(success=False, error="python-docx not installed")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))


class ExcelProcessor:
    def read(self, path: Union[str, Path], sheet: Optional[str] = None) -> DocumentResult:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            sheet_names = wb.sheetnames

            if sheet:
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(cell) if cell is not None else "" for cell in row])

            content = "\n".join([",".join(row) for row in rows])

            return DocumentResult(
                success=True,
                content=content,
                metadata={
                    "path": str(path),
                    "sheets": sheet_names,
                    "active_sheet": sheet,
                    "rows": len(rows),
                },
            )
        except ImportError:
            return DocumentResult(success=False, error="openpyxl not installed. Run: pip install openpyxl")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def get_sheets(self, path: Union[str, Path]) -> List[str]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True)
            return wb.sheetnames
        except Exception:
            return []


class CSVProcessor:
    def read(self, path: Union[str, Path], delimiter: str = ",", limit: int = 10000) -> DocumentResult:
        try:
            import csv
            rows = []
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f, delimiter=delimiter)
                for i, row in enumerate(reader):
                    if i >= limit:
                        break
                    rows.append(delimiter.join(row))
            content = "\n".join(rows)
            return DocumentResult(
                success=True,
                content=content,
                metadata={"path": str(path), "rows": len(rows)},
            )
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def write(self, path: Union[str, Path], content: str, delimiter: str = ",") -> DocumentResult:
        try:
            import csv
            rows = [row.split(delimiter) for row in content.strip().split("\n")]
            with open(path, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f, delimiter=delimiter)
                writer.writerows(rows)
            return DocumentResult(success=True, metadata={"path": str(path), "rows": len(rows)})
        except Exception as e:
            return DocumentResult(success=False, error=str(e))


class JSONProcessor:
    def read(self, path: Union[str, Path]) -> DocumentResult:
        try:
            import json
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            content = json.dumps(data, indent=2)
            return DocumentResult(
                success=True,
                content=content,
                metadata={"path": str(path), "type": type(data).__name__},
            )
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def write(self, path: Union[str, Path], content: str, indent: int = 2) -> DocumentResult:
        try:
            import json
            data = json.loads(content)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=indent)
            return DocumentResult(success=True, metadata={"path": str(path)})
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def validate(self, content: str) -> Dict[str, Any]:
        try:
            import json
            data = json.loads(content)
            return {"valid": True, "type": type(data).__name__}
        except json.JSONDecodeError as e:
            return {"valid": False, "error": str(e)}


class DocumentProcessor:
    def __init__(self):
        self.processors = {
            "txt": TextDocumentProcessor(),
            "md": TextDocumentProcessor(),
            "markdown": TextDocumentProcessor(),
            "json": JSONProcessor(),
            "csv": CSVProcessor(),
            "pdf": PDFProcessor(),
            "docx": WordProcessor(),
            "xlsx": ExcelProcessor(),
        }

    def read(self, path: Union[str, Path]) -> DocumentResult:
        path = Path(path)
        ext = path.suffix.lstrip(".").lower()

        if ext in self.processors:
            return self.processors[ext].read(path)

        return TextDocumentProcessor().read(path)

    def write(self, path: Union[str, Path], content: str) -> DocumentResult:
        path = Path(path)
        ext = path.suffix.lstrip(".").lower()

        if ext in self.processors and hasattr(self.processors[ext], "write"):
            return self.processors[ext].write(path, content)

        return TextDocumentProcessor().write(path, content)

    def get_supported_formats(self) -> List[str]:
        return list(self.processors.keys())
